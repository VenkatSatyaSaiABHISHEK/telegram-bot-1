from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, MessageHandler, ContextTypes, filters
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
import re, os

# Store user text for format selection
user_text_cache = {}

# --- Text Parsing ---
def parse_text(text):
    code_blocks = re.findall(r'```(.*?)```', text, re.DOTALL)
    non_code_parts = re.split(r'```.*?```', text, flags=re.DOTALL)
    return non_code_parts, code_blocks

# --- DOCX Generation ---
def style_code_block(paragraph):
    run = paragraph.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    paragraph._p.get_or_add_pPr().append(shading_elm)

def create_docx(text):
    doc = Document()
    non_code, code_blocks = parse_text(text)
    for i, part in enumerate(non_code):
        if part.strip():
            doc.add_paragraph(part.strip())
        if i < len(code_blocks):
            para = doc.add_paragraph()
            para.add_run(code_blocks[i].strip())
            style_code_block(para)
    file_path = "output.docx"
    doc.save(file_path)
    return file_path

# --- PDF Generation ---
def create_pdf(text):
    filename = "output.pdf"
    doc = SimpleDocTemplate(filename, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    code_style = ParagraphStyle(name='CodeBlock', fontName='Courier', fontSize=9, backColor=colors.whitesmoke, leftIndent=10, borderPadding=5, leading=12, alignment=TA_LEFT)
    elements = []
    non_code, code_blocks = parse_text(text)
    for i, part in enumerate(non_code):
        if part.strip():
            elements.append(Paragraph(part.strip().replace("\n", "<br/>"), normal_style))
            elements.append(Spacer(1, 12))
        if i < len(code_blocks):
            code_text = code_blocks[i].strip().replace(" ", "&nbsp;").replace("\n", "<br/>")
            elements.append(Paragraph(code_text, code_style))
            elements.append(Spacer(1, 12))
    doc.build(elements)
    return filename

# --- Telegram Bot Handlers ---
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📝 Convert Text", callback_data='convert')],
        [InlineKeyboardButton("❓ Help", callback_data='help')],
        [InlineKeyboardButton("⚙️ Settings", callback_data='settings')]
    ]
    await update.message.reply_text(
        "👋 *Welcome to Text Converter Bot!*\n\n"
        "Send any text (with or without code) and choose between PDF or DOCX output.",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data == 'convert':
        await query.edit_message_text(
            "✍️ Send me the text you'd like to convert.\n\n"
            "📌 Tip: Wrap code in triple backticks like:\n```python\nprint('Hello')\n```"
        )
    elif query.data == 'help':
        await query.edit_message_text(
            "ℹ️ *Help*\n\n• Send normal text or code\n• Code blocks: use triple backticks\n• Choose PDF or DOCX\n",
            parse_mode="Markdown"
        )
    elif query.data == 'settings':
        await query.edit_message_text("⚙️ Settings coming soon...")

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_text = update.message.text
    user_id = update.message.from_user.id
    user_text_cache[user_id] = user_text

    keyboard = [
        [InlineKeyboardButton("📄 Convert to DOCX", callback_data='docx')],
        [InlineKeyboardButton("🧾 Convert to PDF", callback_data='pdf')],
        [InlineKeyboardButton("🏠 Home", callback_data='home')]
    ]
    await update.message.reply_text(
        "✅ Text received!\nChoose output format:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def format_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = query.from_user.id
    await query.answer()

    text = user_text_cache.get(user_id)
    if not text:
        await query.edit_message_text("⚠️ Please send text first.")
        return

    if query.data == 'docx':
        file_path = create_docx(text)
        await query.message.reply_document(open(file_path, 'rb'))
    elif query.data == 'pdf':
        file_path = create_pdf(text)
        await query.message.reply_document(open(file_path, 'rb'))

    os.remove(file_path)
    await query.message.reply_text("✅ Done! Want to convert more?", reply_markup=InlineKeyboardMarkup([
        [InlineKeyboardButton("🔁 Convert Another", callback_data='convert')],
        [InlineKeyboardButton("🏠 Home", callback_data='home')]
    ]))

async def handle_home(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await start(query, context)

# --- Main ---
def main():
    TOKEN = os.getenv("BOT_TOKEN")
    if not TOKEN:
        print("❌ Error: BOT_TOKEN not found in environment variables.")
        return

    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CallbackQueryHandler(button_handler, pattern="^(convert|help|settings)$"))
    app.add_handler(CallbackQueryHandler(format_selection, pattern="^(docx|pdf)$"))
    app.add_handler(CallbackQueryHandler(handle_home, pattern="home"))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    print("✅ Bot is running...")
    app.run_polling()

if __name__ == '__main__':
    main()
